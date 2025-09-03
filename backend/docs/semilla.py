from docx import Document
from datetime import datetime

doc = Document()
doc.add_heading("Manual: Estrategia Inicial de Roles y Usuarios (Sin created_by en Roles)", level=1)
doc.add_paragraph("Generado: " + datetime.utcnow().isoformat() + "Z")

secciones = [
("1. Introducción",
 "Necesitas crear usuarios que referencian un rol (role_id), pero los roles normalmente son creados por un usuario. "
 "Al inicio no existen usuarios, lo que genera un dilema circular. Estrategia: simplificar fase inicial omitiendo created_by en roles "
 "y posponer la auditoría hasta que haya base de usuarios."),
("2. Problema Inicial",
 "Circularidad: para crear roles necesitas un creador (usuario), pero para crear el primer usuario quieres asignar un rol. "
 "Resultado: bloqueo si impones FK created_by desde el inicio."),
("3. Criterios de Diseño",
 "- Arrancar rápido sin datos ficticios.\n- Mantener posibilidad de auditoría futura.\n- Minimizar migraciones disruptivas.\n- Evitar hack de usuario fantasma si no es aún necesario."),
("4. Estrategia Elegida (Fase 1 sin created_by)",
 "Definir tabla roles sin columna created_by. Sembrar (seed) roles base mediante migración. Crear primer usuario asignándole rol explícitamente. "
 "Más adelante (cuando haya usuarios administradores) se agrega created_by con migración de compatibilidad."),
("5. Flujo Bootstrap",
 "1) Migración crea tabla roles y siembra roles base (SUPERADMIN, ADMIN, etc.).\n"
 "2) Endpoint/CLI crea primer usuario y le asigna rol SUPERADMIN.\n"
 "3) A partir del segundo usuario, flujo normal de registro/asignación.\n"
 "4) Auditoría (created_by) se incorpora en Fase 2."),
("6. Pasos Técnicos Detallados",
 "Paso A: Crear tabla roles mínima.\nPaso B: Sembrar roles base en migración idempotente.\nPaso C: Modelo User con campo role_id (FK nullable inicialmente si quieres crear usuario antes de asignar).\n"
 "Paso D: Servicio para crear primer usuario (comando administrativo).\nPaso E: Endpoint de gestión de roles (crear/editar) restringido a usuarios con rol privilegiado."),
("7. Migraciones (Fase Inicial)",
 "Migración 1: crear roles.\nMigración 2 (opcional, mismo bloque): seed roles.\nSin created_by aún."),
("8. Añadir created_by Más Adelante (Fase 2)",
 "1) Añadir columna created_by NULLABLE.\n2) Rellenar created_by para roles existentes (ej: asignar al primer SUPERADMIN creado o dejar NULL si es permitido).\n"
 "3) Opcional: volverla NOT NULL + crear usuario system o mantener nullable sólo para roles seed.\n4) Actualizar servicios para set created_by = current_user.id."),
("9. Alternativa: Usuario system (Resumen)",
 "Crear usuario interno con UUID fijo y usarlo como created_by de los roles seed. Permite columna NOT NULL desde el inicio, pero añade complejidad y controles."),
("10. Comparación Rápida",
 "Sin created_by (fase 1): +simple, +rápido, -auditoría inicial.\nUsuario system: +auditoría, +consistencia, +complejidad.\nNullable + system luego: equilibrio."),
("11. Checklist de Implementación",
 "[ ] Crear migración roles.\n[ ] Seed roles base.\n[ ] Agregar role_id a users.\n[ ] Comando crear primer usuario.\n[ ] Validar restricciones de asignación.\n[ ] Planificada migración futura created_by."),
("12. Riesgos y Mitigaciones",
 "Riesgo: Pérdida de autoría inicial -> Documentar seed.\nRiesgo: Scripts asignen roles inexistentes -> FK + validación.\n"
 "Riesgo: Olvidar migración futura -> Registrar tarea en backlog técnica."),
("13. Plan de Evolución",
 "Fase 1 (MVP): Sin created_by.\nFase 2: Añadir created_by nullable + rellenar.\nFase 3: Auditoría completa (tabla roles_audit, triggers o capa servicio)."),
("14. Apéndice: Snippets",
 "Incluye ejemplos de modelo, migración y comando inicial.")
]

for titulo, texto in secciones:
    doc.add_heading(titulo, level=2)
    for linea in texto.split("\n"):
        doc.add_paragraph(linea)

doc.add_page_break()
doc.add_heading("Snippets Técnicos", level=1)

doc.add_heading("Modelo Role (Fase 1)", level=2)
doc.add_paragraph(
"""class Role(Base):
    __tablename__ = "roles"
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    name = Column(String(64), unique=True, nullable=False)
    description = Column(Text)
    created_at = Column(DateTime(timezone=True), server_default=func.now())"""
)

doc.add_heading("Modelo User (role_id)", level=2)
doc.add_paragraph(
"""class User(Base):
    __tablename__ = "users"
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    username = Column(String(64), unique=True, nullable=False)
    role_id = Column(UUID(as_uuid=True), ForeignKey("roles.id"), nullable=True)
    role = relationship("Role")"""
)

doc.add_heading("Migración Roles + Seed (Ejemplo)", level=2)
doc.add_paragraph(
"""def upgrade():
    op.create_table(
        "roles",
        sa.Column("id", sa.dialects.postgresql.UUID(as_uuid=True), primary_key=True),
        sa.Column("name", sa.String(64), nullable=False, unique=True),
        sa.Column("description", sa.Text()),
        sa.Column("created_at", sa.DateTime(timezone=True), server_default=sa.text("now()"))
    )
    conn = op.get_bind()
    conn.execute(sa.text(\"\"\"INSERT INTO roles (id,name,description) VALUES
        (gen_random_uuid(),'SUPERADMIN','Rol con todos los permisos'),
        (gen_random_uuid(),'ADMIN','Administración general')
        ON CONFLICT DO NOTHING;\"\"\"))"""
)

doc.add_heading("Migración Fase 2 (Añadir created_by)", level=2)
doc.add_paragraph(
"""def upgrade():
    op.add_column('roles', sa.Column('created_by', sa.dialects.postgresql.UUID(as_uuid=True), sa.ForeignKey('users.id'), nullable=True))
    # Rellenar si decides (ejemplo): UPDATE roles SET created_by = <primer_superadmin_id>;
    # Luego opcional: op.alter_column('roles','created_by', nullable=False)"""
)

doc.save("docs/manual_roles_usuarios.docx")
print("Archivo generado en docs/manual_roles_usuarios.docx")